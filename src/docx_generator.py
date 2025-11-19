"""
Professional DOCX Generator - Create beautifully formatted Word documents
Formatting: Times New Roman, 12pt, 1.5 spacing, A4, proper headers, bold/italic
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re
import os

class ProfessionalDocxGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_document_formatting()
    
    def _setup_document_formatting(self):
        """Set up professional document formatting."""
        # Page setup - A4
        section = self.doc.sections[0]
        section.page_height = Inches(11.69)  # A4 height
        section.page_width = Inches(8.27)    # A4 width
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Default style - Times New Roman, 12pt, 1.5 spacing
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)
        
        # Heading 1 style
        h1_style = self.doc.styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)
        
        # Heading 2 style
        h2_style = self.doc.styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(16)
        h2_font.bold = True
        
        # Heading 3 style
        h3_style = self.doc.styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(14)
        h3_font.bold = True
        
        # Heading 4 style
        h4_style = self.doc.styles['Heading 4']
        h4_font = h4_style.font
        h4_font.name = 'Times New Roman'
        h4_font.size = Pt(12)
        h4_font.bold = True
    
    def convert_markdown_to_docx(self, md_file: str, output_file: str = None):
        """
        Convert markdown file to professionally formatted DOCX.
        
        Args:
            md_file: Path to markdown file
            output_file: Output DOCX path (optional)
        """
        if output_file is None:
            output_file = md_file.replace('.md', '.docx')
        
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Process line by line
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Skip metadata lines
            if line.startswith('*Generation') or line.startswith('*Total words'):
                i += 1
                continue
            
            # Horizontal rules
            if line.strip() == '---':
                self.doc.add_paragraph()  # Just add space
                i += 1
                continue
            
            # Headings
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1)
                i += 1
                continue
            
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], level=2)
                i += 1
                continue
            
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], level=3)
                i += 1
                continue
            
            elif line.startswith('#### '):
                self.doc.add_heading(line[5:], level=4)
                i += 1
                continue
            
            # Code blocks (ASCII diagrams)
            elif line.strip().startswith('```'):
                # Collect code block
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                # Add as monospace paragraph
                if code_lines:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                
                i += 1
                continue
            
            # Regular paragraphs
            elif line.strip():
                # Process inline formatting
                p = self.doc.add_paragraph()
                self._add_formatted_text(p, line)
                i += 1
                continue
            
            # Empty lines
            else:
                i += 1
                continue
        
        # Save
        self.doc.save(output_file)
        return output_file
    
    def _add_formatted_text(self, paragraph, text):
        """Add text with bold, italic, and other formatting."""
        # Pattern for bold, italic, code
        # **bold**, *italic*, `code`
        
        parts = []
        current = ""
        i = 0
        
        while i < len(text):
            # Bold **text**
            if text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing **
                j = text.find('**', i+2)
                if j != -1:
                    parts.append(('bold', text[i+2:j]))
                    i = j + 2
                else:
                    current += text[i]
                    i += 1
            
            # Italic *text* (but not **)
            elif text[i] == '*' and (i == 0 or text[i-1] != '*') and (i+1 >= len(text) or text[i+1] != '*'):
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing *
                j = i + 1
                while j < len(text) and text[j] != '*':
                    j += 1
                
                if j < len(text):
                    parts.append(('italic', text[i+1:j]))
                    i = j + 1
                else:
                    current += text[i]
                    i += 1
            
            # Code `text`
            elif text[i] == '`':
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing `
                j = text.find('`', i+1)
                if j != -1:
                    parts.append(('code', text[i+1:j]))
                    i = j + 1
                else:
                    current += text[i]
                    i += 1
            
            else:
                current += text[i]
                i += 1
        
        if current:
            parts.append(('normal', current))
        
        # Add runs with formatting
        for fmt, txt in parts:
            run = paragraph.add_run(txt)
            
            if fmt == 'bold':
                run.bold = True
            elif fmt == 'italic':
                run.italic = True
            elif fmt == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(11)

def auto_convert_to_docx(md_file: str) -> str:
    """
    Automatically convert markdown to DOCX with professional formatting.
    
    Args:
        md_file: Path to markdown file
    
    Returns:
        Path to generated DOCX file
    """
    try:
        generator = ProfessionalDocxGenerator()
        docx_file = generator.convert_markdown_to_docx(md_file)
        print(f"✅ Created DOCX: {docx_file}")
        return docx_file
    except Exception as e:
        print(f"❌ DOCX conversion failed: {e}")
        return None

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        auto_convert_to_docx(sys.argv[1])
    else:
        print("Usage: python3 -m src.docx_generator <markdown_file>")
