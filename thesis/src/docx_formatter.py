"""
DOCX Formatter for PhD Thesis
Converts markdown to DOCX with strict thesis formatting standards
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re


class ThesisDocxFormatter:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for thesis formatting"""
        styles = self.doc.styles
        
        # Normal style - Body text
        normal = styles['Normal']
        normal_font = normal.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        normal_font.color.rgb = RGBColor(0, 0, 0)
        
        normal_para = normal.paragraph_format
        normal_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # H1 - Chapter titles (Centered, UPPERCASE, Bold, 14pt)
        try:
            h1 = styles['Heading 1']
        except KeyError:
            h1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        
        h1_font = h1.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(14)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)
        
        h1_para = h1.paragraph_format
        h1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_para.space_before = Pt(24)
        h1_para.space_after = Pt(12)
        h1_para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # H2 - Section titles (Bold, 13pt)
        try:
            h2 = styles['Heading 2']
        except KeyError:
            h2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        
        h2_font = h2.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(13)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
        
        h2_para = h2.paragraph_format
        h2_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        h2_para.space_before = Pt(18)
        h2_para.space_after = Pt(6)
        h2_para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # H3 - Subsection titles (Bold, Italic, 12pt)
        try:
            h3 = styles['Heading 3']
        except KeyError:
            h3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        
        h3_font = h3.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.italic = True
        h3_font.color.rgb = RGBColor(0, 0, 0)
        
        h3_para = h3.paragraph_format
        h3_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        h3_para.space_before = Pt(12)
        h3_para.space_after = Pt(6)
        h3_para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    def markdown_to_docx(self, markdown_file, output_file):
        """Convert markdown file to formatted DOCX"""
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # H1 - Chapter titles (# CHAPTER)
            if line.startswith('# CHAPTER'):
                text = line[2:].strip().upper()
                p = self.doc.add_paragraph(text, style='Heading 1')
            
            # H2 - Section titles (##)
            elif line.startswith('## '):
                text = line[3:].strip()
                p = self.doc.add_paragraph(text, style='Heading 2')
            
            # H3 - Subsection titles (###)
            elif line.startswith('### '):
                text = line[4:].strip()
                p = self.doc.add_paragraph(text, style='Heading 3')
            
            # H4 - Sub-subsection (####)
            elif line.startswith('#### '):
                text = line[5:].strip()
                p = self.doc.add_paragraph(text, style='Heading 3')  # Use H3 style
            
            # Tables
            elif line.startswith('|') and '|' in line:
                table_lines = [line]
                i += 1
                while i < len(lines) and lines[i].startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                i -= 1  # Back up one
                self._add_table(table_lines)
            
            # Code blocks (for ASCII charts)
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    p = self.doc.add_paragraph('\n'.join(code_lines))
                    p.paragraph_format.left_indent = Inches(0.5)
            
            # Bold text (**text**)
            elif '**' in line:
                self._add_formatted_paragraph(line)
            
            # Italic text (*text*)
            elif '*' in line and not line.startswith('*Source:'):
                self._add_formatted_paragraph(line)
            
            # Regular paragraph
            elif line.strip():
                p = self.doc.add_paragraph(line)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Empty line
            else:
                pass  # Skip empty lines
            
            i += 1
        
        self.doc.save(output_file)
        print(f"✅ DOCX saved: {output_file}")
    
    def _add_table(self, table_lines):
        """Add a formatted table"""
        # Parse table
        rows = []
        for line in table_lines:
            if '---' in line:  # Skip separator line
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
            rows.append(cells)
        
        if not rows:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        # Fill table
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
                
                # Format cell
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Bold header row
                    if i == 0:
                        for run in paragraph.runs:
                            run.font.bold = True
                    
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def _add_formatted_paragraph(self, text):
        """Add paragraph with bold/italic formatting"""
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Simple regex for **bold** and *italic*
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                run = p.add_run(part)
            
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
